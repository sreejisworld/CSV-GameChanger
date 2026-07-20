"""
Word Template Injection for the EVOLV platform.

Accepts a user-uploaded .docx template containing {{PLACEHOLDER}}
markers, injects generated validation data (UR/FR, test scripts,
metadata), and returns the populated document as bytes while
preserving the template's formatting, headers, and footers.

:requirement: URS-18.1 - Generate combined Validation Report.
"""

import copy
import io
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table


# ---------------------------------------------------------------
# Placeholder pattern
# ---------------------------------------------------------------
_PH_RE = re.compile(r"\{\{([A-Z_]+)\}\}")


# ---------------------------------------------------------------
# Public API
# ---------------------------------------------------------------
def inject_template(
    template_bytes: bytes,
    ur_fr: Dict[str, Any],
    test_script: Optional[Dict[str, Any]] = None,
    signer_name: str = "",
) -> bytes:
    """
    Inject UR/FR and test-script data into a .docx template.

    Unknown placeholders are left as-is so the user can keep
    static markers they do not want replaced.

    :param template_bytes: Raw bytes of the uploaded .docx.
    :param ur_fr: UR/FR dict from RequirementArchitect.
    :param test_script: Test-script dict from DeltaAgent
        (optional).
    :param signer_name: Full name of the approver.
    :return: Populated .docx as bytes.
    :requirement: URS-18.1 - Generate combined Validation
        Report.
    """
    doc = Document(io.BytesIO(template_bytes))
    ph_map = _build_placeholder_map(
        ur_fr, test_script, signer_name,
    )

    # Phase 1 — table placeholders in body paragraphs
    _replace_table_placeholders(doc, ur_fr, test_script)

    # Phase 2 — text placeholders in body paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, ph_map)

    # Phase 3 — text placeholders inside existing tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, ph_map)

    # Phase 4 — headers and footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, ph_map)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, ph_map)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------
# Placeholder map builder
# ---------------------------------------------------------------
def _build_placeholder_map(
    ur_fr: Dict[str, Any],
    test_script: Optional[Dict[str, Any]],
    signer_name: str,
) -> Dict[str, str]:
    """
    Build a flat ``{placeholder: value}`` map from UR/FR,
    test-script, and signer data.

    :param ur_fr: UR/FR dict.
    :param test_script: Test-script dict (may be None).
    :param signer_name: Approver name.
    :return: Mapping of placeholder key to replacement text.
    """
    u_req = ur_fr.get("user_requirement", {})
    ctx = ur_fr.get("additional_context") or {}

    ph: Dict[str, str] = {
        "URS_ID": str(ur_fr.get("urs_id", "")),
        "REQUIREMENT_SUMMARY": str(
            ur_fr.get("requirement_summary", "")
        ),
        "CATEGORY": str(ur_fr.get("category", "")),
        "RISK_ASSESSMENT": str(
            u_req.get("risk_assessment", "")
        ),
        "IMPLEMENTATION_METHOD": str(
            u_req.get("implementation_method", "")
        ),
        "RISK_LEVEL": str(u_req.get("risk_level", "")),
        "TEST_STRATEGY": str(
            u_req.get("test_strategy", "")
        ),
        "UR_STATEMENT": str(u_req.get("statement", "")),
        "SYSTEM_DESCRIPTION": str(
            ctx.get("system_description", "")
        ),
        "WORKSHOP_NOTES": str(
            ctx.get("workshop_notes", "")
        ),
        "ROLES_AND_PERMISSIONS": str(
            ctx.get("roles_and_permissions", "")
        ),
        "ASSUMPTIONS": "\n".join(
            str(a)
            for a in ur_fr.get(
                "assumptions_and_dependencies", []
            )
        ),
        "COMPLIANCE_NOTES": "\n".join(
            str(n)
            for n in ur_fr.get("compliance_notes", [])
        ),
        "GENERATED_DATE": datetime.now(
            timezone.utc
        ).isoformat(),
        "SIGNER_NAME": signer_name,
    }
    return ph


# ---------------------------------------------------------------
# Text replacement in a single paragraph (run-aware)
# ---------------------------------------------------------------
def _replace_in_paragraph(
    paragraph: Any,
    placeholder_map: Dict[str, str],
) -> None:
    """
    Replace ``{{KEY}}`` tokens in *paragraph* while handling
    Word's run-splitting.

    When a placeholder like ``{{URS_ID}}`` is split across
    multiple runs (e.g. ``{{``, ``URS_ID``, ``}}``), the runs
    are consolidated so the replacement can succeed.

    :param paragraph: A ``docx.text.paragraph.Paragraph``.
    :param placeholder_map: ``{KEY: value}`` mapping.
    """
    full_text = paragraph.text
    if "{{" not in full_text:
        return

    # Check if any known placeholder is present
    has_match = False
    for key in placeholder_map:
        token = "{{" + key + "}}"
        if token in full_text:
            has_match = True
            break

    if not has_match:
        return

    # Consolidate runs: rebuild text, apply replacements,
    # then put the result into the first run and clear the rest.
    for key, value in placeholder_map.items():
        token = "{{" + key + "}}"
        full_text = full_text.replace(token, value)

    runs = paragraph.runs
    if not runs:
        return

    runs[0].text = full_text
    for run in runs[1:]:
        run.text = ""


# ---------------------------------------------------------------
# Table placeholder handling
# ---------------------------------------------------------------
_REQ_TABLE_COLS = [
    "FR ID", "Parent UR", "Statement",
    "Acceptance Criteria",
]

_TEST_TABLE_COLS = [
    "Type", "#", "Title", "Instruction",
    "Expected Result", "Case", "Ref",
]


def _replace_table_placeholders(
    doc: Any,
    ur_fr: Dict[str, Any],
    test_script: Optional[Dict[str, Any]],
) -> None:
    """
    Scan body paragraphs for ``{{REQUIREMENTS_TABLE}}`` and
    ``{{TEST_STEPS_TABLE}}``.  For each match, insert a real
    docx Table and remove the placeholder paragraph.

    :param doc: The ``Document`` object.
    :param ur_fr: UR/FR dict.
    :param test_script: Test-script dict (may be None).
    """
    # Collect paragraphs to process (iterate over a snapshot)
    to_replace: List[Tuple[Any, str]] = []
    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        if txt == "{{REQUIREMENTS_TABLE}}":
            to_replace.append(
                (paragraph, "REQUIREMENTS_TABLE")
            )
        elif txt == "{{TEST_STEPS_TABLE}}":
            to_replace.append(
                (paragraph, "TEST_STEPS_TABLE")
            )

    for paragraph, kind in to_replace:
        if kind == "REQUIREMENTS_TABLE":
            rows = _build_req_rows(ur_fr)
            _insert_table_after_paragraph(
                paragraph, doc, _REQ_TABLE_COLS, rows,
            )
        else:
            rows = _build_test_rows(test_script)
            _insert_table_after_paragraph(
                paragraph, doc, _TEST_TABLE_COLS, rows,
            )


def _build_req_rows(
    ur_fr: Dict[str, Any],
) -> List[List[str]]:
    """
    Build requirement-table data rows from functional
    requirements.

    :param ur_fr: UR/FR dict.
    :return: List of row-value lists.
    """
    rows: List[List[str]] = []
    for fr in ur_fr.get("functional_requirements", []):
        criteria = fr.get("acceptance_criteria", [])
        criteria_text = "\n".join(
            str(c) for c in criteria
        )
        rows.append([
            str(fr.get("fr_id", "")),
            str(fr.get("parent_ur_id", "")),
            str(fr.get("statement", "")),
            criteria_text,
        ])
    return rows


def _build_test_rows(
    test_script: Optional[Dict[str, Any]],
) -> List[List[str]]:
    """
    Build test-step-table data rows from a test script.

    :param test_script: Test-script dict (may be None).
    :return: List of row-value lists (empty if no script).
    """
    if test_script is None:
        return []
    rows: List[List[str]] = []
    for s in test_script.get("steps", []):
        rows.append([
            str(s.get("step_type", "")),
            str(s.get("step_number", "")),
            str(s.get("step_title", "")),
            str(s.get("step_instruction", "")),
            str(s.get("expected_result", "")),
            str(s.get("test_case_type", "")),
            str(s.get("requirement_reference", "")),
        ])
    return rows


def _insert_table_after_paragraph(
    paragraph: Any,
    doc: Any,
    columns: List[str],
    rows: List[List[str]],
) -> None:
    """
    Insert a docx Table immediately after *paragraph* in the
    document body, then remove the placeholder paragraph.

    :param paragraph: The placeholder paragraph to replace.
    :param doc: The ``Document`` object.
    :param columns: Column header labels.
    :param rows: Data rows (list of lists).
    """
    num_rows = 1 + len(rows)  # header + data
    num_cols = len(columns)
    table = doc.add_table(
        rows=num_rows, cols=num_cols,
    )
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    # Header row
    for i, col_name in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = col_name
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # Move the table XML element right after the paragraph
    p_element = paragraph._element
    p_element.addnext(table._tbl)

    # Remove the placeholder paragraph
    p_parent = p_element.getparent()
    if p_parent is not None:
        p_parent.remove(p_element)
