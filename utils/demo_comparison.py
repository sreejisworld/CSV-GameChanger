"""
Demo Comparison utilities for the EVOLV Engine.

Provides deterministic requirement analysis, rewriting, and
gap detection for the sales demo comparison page.  No
Pinecone or OpenAI calls are needed.

:requirement: URS-19.1 - Detect ambiguous language in
    human-written requirements.
:requirement: URS-19.2 - Detect missing regulatory controls.
:requirement: URS-19.3 - Rewrite requirements to audit-ready
    format.
:requirement: URS-19.4 - Evaluate and score requirement
    quality.
:requirement: URS-19.5 - Inject AI output into .docx
    templates.
"""

from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, List, Tuple


# ---------------------------------------------------------------
# Constants
# ---------------------------------------------------------------

AMBIGUITY_WORDS: List[str] = [
    "should",
    "fast",
    "easy",
    "all",
    "quickly",
    "user-friendly",
    "intuitive",
    "robust",
    "seamless",
    "appropriate",
    "adequate",
    "timely",
    "efficient",
    "flexible",
    "as needed",
    "etc",
    "reasonable",
    "minimal",
    "proper",
    "sufficient",
]

REGULATORY_CONTROLS: Dict[str, Dict[str, Any]] = {
    "Audit Trail": {
        "keywords": [
            "audit trail",
            "audit log",
            "traceability",
            "21 cfr part 11",
        ],
        "injection": (
            "The system shall maintain an append-only, "
            "tamper-evident audit trail recording all data "
            "creation, modification, and deletion events "
            "per 21 CFR Part 11."
        ),
    },
    "Electronic Signatures": {
        "keywords": [
            "electronic signature",
            "e-signature",
            "digital signature",
            "two-component",
        ],
        "injection": (
            "The system shall support two-component "
            "electronic signatures (user ID + password) "
            "with configurable signature meanings per "
            "21 CFR Part 11."
        ),
    },
    "Access Controls": {
        "keywords": [
            "access control",
            "role-based",
            "rbac",
            "permission",
            "privilege",
            "authentication",
        ],
        "injection": (
            "The system shall enforce role-based access "
            "controls restricting functionality by user "
            "role with documented privilege matrices."
        ),
    },
    "Data Integrity": {
        "keywords": [
            "data integrity",
            "alcoa",
            "checksum",
            "backup",
            "attributable",
            "legible",
            "contemporaneous",
        ],
        "injection": (
            "The system shall ensure data integrity "
            "following ALCOA+ principles: data shall be "
            "attributable, legible, contemporaneous, "
            "original, and accurate."
        ),
    },
}

COST_PER_POOR_REQUIREMENT: int = 5000


# ---------------------------------------------------------------
# Public functions
# ---------------------------------------------------------------

def detect_ambiguities(human_text: str) -> List[str]:
    """Scan *human_text* for vague / ambiguous words.

    :param human_text: Original requirement text.
    :return: List of ambiguity words found.
    :requirement: URS-19.1 - Detect ambiguous language.
    """
    lower = human_text.lower()
    found: List[str] = []
    for word in AMBIGUITY_WORDS:
        if word in lower:
            found.append(word)
    return found


def detect_regulatory_gaps(
    human_text: str,
) -> List[Dict[str, str]]:
    """Check *human_text* for missing regulatory controls.

    Each REGULATORY_CONTROLS category is checked; if no
    keyword matches, a gap is reported with the category
    name and recommended injection clause.

    :param human_text: Original requirement text.
    :return: List of gap dicts with *category* and
        *injection* keys.
    :requirement: URS-19.2 - Detect missing regulatory
        controls.
    """
    lower = human_text.lower()
    gaps: List[Dict[str, str]] = []
    for category, ctrl in REGULATORY_CONTROLS.items():
        if not any(kw in lower for kw in ctrl["keywords"]):
            gaps.append({
                "category": category,
                "injection": ctrl["injection"],
            })
    return gaps


def rewrite_requirement(
    human_text: str,
) -> Tuple[str, str]:
    """Rewrite a human requirement to audit-ready format.

    Uses the RequirementArchitect's deterministic helpers
    (``_format_requirement_statement`` and
    ``_determine_criticality``) without any Pinecone or
    OpenAI calls.

    :param human_text: Original requirement text.
    :return: Tuple of (ai_rewritten_text, criticality).
    :requirement: URS-19.3 - Rewrite requirements to
        audit-ready format.
    """
    from Agents.requirement_architect import (
        RequirementArchitect,
    )

    # Build a lightweight instance that skips Pinecone
    arch = RequirementArchitect.__new__(
        RequirementArchitect,
    )
    arch._urs_counter = 0
    arch._template = None
    arch._vector_search_available = False
    arch._pinecone_api_key = None
    arch._openai_api_key = None
    arch._index_name = ""

    # Pre-clean common prefixes that the architect
    # doesn't strip itself
    cleaned = human_text.strip()
    _extra_prefixes = [
        "the system should",
        "the application should",
        "the software should",
        "users need to be able to",
        "users should be able to",
    ]
    _lower = cleaned.lower()
    for _pf in _extra_prefixes:
        if _lower.startswith(_pf):
            cleaned = cleaned[len(_pf):].strip()
            break

    # Format to "The system shall ..." style
    formatted = arch._format_requirement_statement(
        cleaned,
    )

    # Classify criticality (keyword-based, no search)
    criticality = arch._determine_criticality(
        human_text, search_results=[],
    )
    crit_str = criticality.value  # "High" / "Medium" / "Low"

    # Append up to 2 missing regulatory clauses
    gaps = detect_regulatory_gaps(human_text)
    injections = [g["injection"] for g in gaps[:2]]
    if injections:
        formatted = (
            formatted
            + " Additionally, "
            + " ".join(injections)
        )

    return formatted, crit_str


def evaluate_requirements(
    human_text: str,
    ai_text: str,
    criticality: str,
) -> Dict[str, Any]:
    """Score a human vs AI requirement pair.

    :param human_text: Original requirement text.
    :param ai_text: AI-rewritten requirement text.
    :param criticality: Criticality classification string.
    :return: Dict with evaluation metrics.
    :requirement: URS-19.4 - Evaluate and score requirement
        quality.
    """
    ambiguities = detect_ambiguities(human_text)
    gaps = detect_regulatory_gaps(human_text)

    risk_bullets: List[str] = []
    if ambiguities:
        risk_bullets.append(
            f"Ambiguous language: "
            f"{', '.join(ambiguities)}"
        )
    for g in gaps:
        risk_bullets.append(
            f"Missing {g['category']} controls"
        )

    issue_count = len(ambiguities) + len(gaps)
    cost_of_error = issue_count * COST_PER_POOR_REQUIREMENT

    return {
        "human_text": human_text,
        "ai_text": ai_text,
        "criticality": criticality,
        "ambiguities": ambiguities,
        "regulatory_gaps": gaps,
        "risk_bullets": risk_bullets,
        "issue_count": issue_count,
        "cost_of_error": cost_of_error,
    }


def inject_into_docx(
    template_bytes: bytes,
    system_desc: str,
    comparisons: List[Dict[str, Any]],
) -> bytes:
    """Inject comparison results into a .docx template.

    If the template contains placeholders
    (``{{SYSTEM_DESCRIPTION}}``, ``{{REQUIREMENT_1}}``,
    etc.) they are replaced.  Otherwise content is appended
    at the end of the document.

    :param template_bytes: Uploaded .docx as bytes.
    :param system_desc: System description text.
    :param comparisons: List of evaluation dicts from
        :func:`evaluate_requirements`.
    :return: Modified .docx as bytes.
    :requirement: URS-19.5 - Inject AI output into .docx
        templates.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document(BytesIO(template_bytes))

    # ---- Try placeholder replacement first ----
    placeholders_found = False
    for para in doc.paragraphs:
        if "{{SYSTEM_DESCRIPTION}}" in para.text:
            para.text = para.text.replace(
                "{{SYSTEM_DESCRIPTION}}", system_desc,
            )
            placeholders_found = True
        for i, comp in enumerate(comparisons, 1):
            tag = f"{{{{REQUIREMENT_{i}}}}}"
            if tag in para.text:
                para.text = para.text.replace(
                    tag, comp["ai_text"],
                )
                placeholders_found = True

    if not placeholders_found:
        # ---- Append mode ----
        doc.add_heading("System Description", level=2)
        doc.add_paragraph(system_desc)

        doc.add_heading(
            "AI-Generated Requirements", level=2,
        )
        for i, comp in enumerate(comparisons, 1):
            p = doc.add_paragraph()

            # Requirement number
            run = p.add_run(f"Requirement {i}:  ")
            run.bold = True
            run.font.size = Pt(11)

            # AI text
            run = p.add_run(comp["ai_text"])
            run.font.size = Pt(11)

            # Criticality badge (colour-coded bold)
            run = p.add_run(
                f"  [{comp['criticality']}]",
            )
            run.bold = True
            run.font.size = Pt(10)
            crit_lower = comp["criticality"].lower()
            if crit_lower == "high":
                run.font.color.rgb = RGBColor(
                    0xB9, 0x1C, 0x1C,
                )
            elif crit_lower == "medium":
                run.font.color.rgb = RGBColor(
                    0x92, 0x40, 0x0E,
                )
            else:
                run.font.color.rgb = RGBColor(
                    0x06, 0x5F, 0x46,
                )

            # Gap bullets
            for bullet in comp.get("risk_bullets", []):
                bp = doc.add_paragraph(
                    bullet, style="List Bullet",
                )
                for r in bp.runs:
                    r.font.size = Pt(10)

    # Footer
    ts = datetime.now(timezone.utc).strftime(
        "%Y-%m-%d %H:%M UTC",
    )
    fp = doc.add_paragraph()
    run = fp.add_run(
        f"Powered by EVOLV | A WingstarTech Inc. Product | {ts}",
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
