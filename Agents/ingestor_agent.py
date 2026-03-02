"""
Ingestor Agent Module.

Reads vendor .docx and .pdf documents from input/vendor_docs/ and converts
them into structured JSON that the RequirementArchitect can consume.
Also performs GAMP 5 gap analysis against the Pinecone knowledge base.

:requirement: URS-8.1 - System shall ingest vendor documents for processing.
:requirement: URS-9.1 - System shall perform GAMP 5 gap analysis.
"""
import hashlib
import json
import logging
import os
import re
import statistics
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from dataclasses import dataclass, field

from Agents.integrity_manager import (
    log_audit_event as _log_integrity_event,
)

# ── Optional heavy deps ──────────────────────────────────────────────
try:
    import pdfplumber as _pdfplumber
except ImportError:
    _pdfplumber = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn as _docx_qn
    _DOCX_HAS_XML = True
except ImportError:
    DocxDocument = None
    _docx_qn = None
    _DOCX_HAS_XML = False

try:
    from openai import OpenAI as _OpenAI
except ImportError:
    _OpenAI = None

# ── Configuration ─────────────────────────────────────────────────────
_ROOT = Path(__file__).parent.parent
VENDOR_DOCS_DIR = str(_ROOT / "input" / "vendor_docs")
OUTPUT_DIR = str(_ROOT / "output")
MANIFEST_PATH = _ROOT / "output" / "ingest_manifest.json"
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
EVIDENCE_EXCERPT_LENGTH = 200

# GAMP 5 expected lifecycle document categories for gap analysis.
GAMP5_CATEGORIES: List[Dict[str, Any]] = [
    {
        "category": "Intended Use",
        "query": "intended use statement system purpose",
        "keywords": [
            "intended use", "purpose", "scope of use",
            "system purpose", "intended application"
        ]
    },
    {
        "category": "User Requirements (URS)",
        "query": "user requirements specification functional",
        "keywords": [
            "user requirement", "urs", "functional requirement",
            "user need", "business requirement"
        ]
    },
    {
        "category": "Functional Specifications",
        "query": "functional specification design",
        "keywords": [
            "functional specification", "functional spec",
            "system function", "feature", "capability"
        ]
    },
    {
        "category": "Risk Assessment",
        "query": "risk assessment patient safety FMEA",
        "keywords": [
            "risk assessment", "risk analysis", "fmea",
            "hazard", "risk matrix", "patient safety",
            "risk priority"
        ]
    },
    {
        "category": "Design Specification",
        "query": "design specification architecture",
        "keywords": [
            "design specification", "design spec",
            "architecture", "system design", "technical design"
        ]
    },
    {
        "category": "Traceability",
        "query": "traceability matrix requirements testing",
        "keywords": [
            "traceability", "trace matrix",
            "requirements traceability", "rtm"
        ]
    },
    {
        "category": "Testing Strategy",
        "query": "testing strategy IQ OQ PQ validation",
        "keywords": [
            "testing strategy", "test plan", "iq", "oq", "pq",
            "validation protocol", "test case", "test script"
        ]
    },
    {
        "category": "Change Control",
        "query": "change control management procedure",
        "keywords": [
            "change control", "change management",
            "change request", "change procedure"
        ]
    },
    {
        "category": "Data Integrity",
        "query": "data integrity ALCOA electronic records",
        "keywords": [
            "data integrity", "alcoa", "electronic record",
            "audit trail", "21 cfr part 11", "data governance"
        ]
    },
    {
        "category": "Supplier Assessment",
        "query": "supplier assessment vendor audit qualification",
        "keywords": [
            "supplier assessment", "vendor audit",
            "supplier qualification", "vendor qualification",
            "supplier evaluation"
        ]
    },
    {
        "category": "Validation Plan",
        "query": "validation plan approach lifecycle",
        "keywords": [
            "validation plan", "validation approach",
            "validation strategy", "lifecycle",
            "validation lifecycle"
        ]
    },
    {
        "category": "Standard Operating Procedures",
        "query": "SOP standard operating procedure training",
        "keywords": [
            "sop", "standard operating procedure",
            "procedure", "work instruction", "training"
        ]
    },
]

# ── Audit logger ──────────────────────────────────────────────────────
audit_logger = logging.getLogger("csv_engine.audit")

# ── Heading regex ─────────────────────────────────────────────────────
_HEADING_RE = re.compile(
    r'^\d+\.?\d*\.?\s+[A-Z]|^Section\s+\d',
    re.IGNORECASE
)

# ── Limitation patterns ───────────────────────────────────────────────
_LIMITATION_PATTERNS = re.compile(
    r'[^.]*\b(?:shall not|must not|cannot|may not|'
    r'is not permitted|is prohibited|not allowed|'
    r'will not support|out of scope|excluded from)\b[^.]*\.',
    re.IGNORECASE
)


# ── Exceptions ────────────────────────────────────────────────────────
class IngestorError(Exception):
    """
    Base exception for Ingestor Agent errors.

    Error code: CSV-005 - Document ingestion failed.

    :requirement: URS-8.2 - System shall report ingestion errors.
    """
    error_code = "CSV-005"


class UnsupportedFileTypeError(IngestorError):
    """
    Raised when a file type is not supported for ingestion.

    Error code: CSV-006 - Unsupported file type.

    :requirement: URS-8.2 - System shall report ingestion errors.
    """
    error_code = "CSV-006"

    def __init__(self, file_path: str):
        ext = Path(file_path).suffix
        super().__init__(
            f"Unsupported file type '{ext}' for '{file_path}'. "
            f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}"
        )


class DocumentParseError(IngestorError):
    """
    Raised when a document cannot be parsed.

    Error code: CSV-007 - Document parsing failed.

    :requirement: URS-8.2 - System shall report ingestion errors.
    """
    error_code = "CSV-007"

    def __init__(self, file_path: str, reason: str):
        super().__init__(
            f"Failed to parse '{file_path}': {reason}"
        )


# ── Dataclasses ───────────────────────────────────────────────────────
@dataclass
class DocumentSection:
    """
    A section extracted from a vendor document.

    :requirement: URS-8.3 - System shall extract structured sections.
    """
    heading: str
    content: str
    page_number: int
    section_index: int
    section_type: str = "body"  # "body" or "table"

    def to_dict(self) -> Dict[str, Any]:
        """
        Convert section to dictionary.

        :return: Dictionary representation of the section.
        :requirement: URS-8.4 - System shall output JSON format.
        """
        return {
            "heading": self.heading,
            "content": self.content,
            "page_number": self.page_number,
            "section_index": self.section_index,
            "section_type": self.section_type,
        }


@dataclass
class IngestedDocument:
    """
    Structured representation of an ingested vendor document.

    :requirement: URS-8.3 - System shall extract structured sections.
    """
    file_name: str
    file_type: str
    title: str
    ingested_at: str
    total_pages: int
    document_hash: str = ""
    sections: List[DocumentSection] = field(default_factory=list)
    requirements: List[str] = field(default_factory=list)
    limitations: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        """
        Convert ingested document to dictionary.

        :return: Dictionary representation of the document.
        :requirement: URS-8.4 - System shall output JSON format.
        """
        return {
            "file_name": self.file_name,
            "file_type": self.file_type,
            "title": self.title,
            "ingested_at": self.ingested_at,
            "total_pages": self.total_pages,
            "document_hash": self.document_hash,
            "sections": [s.to_dict() for s in self.sections],
            "requirements": self.requirements,
            "limitations": self.limitations,
        }

    def to_json(self) -> str:
        """
        Convert ingested document to JSON string.

        :return: JSON string representation.
        :requirement: URS-8.4 - System shall output JSON format.
        """
        return json.dumps(self.to_dict(), indent=2)


@dataclass
class GapFinding:
    """
    A single finding from GAMP 5 gap analysis.

    :requirement: URS-9.2 - System shall identify gaps per category.
    """
    category: str
    status: str
    vendor_evidence: str
    gamp5_reference: str
    recommendation: str
    similarity_score: float = 0.0
    regulatory_clause_mapping: List[Dict[str, Any]] = field(
        default_factory=list
    )

    def to_dict(self) -> Dict[str, Any]:
        """
        Convert finding to dictionary.

        :return: Dictionary representation of the finding.
        :requirement: URS-9.3 - System shall output gap analysis as JSON.
        """
        return {
            "category": self.category,
            "status": self.status,
            "similarity_score": round(self.similarity_score, 4),
            "vendor_evidence": self.vendor_evidence,
            "gamp5_reference": self.gamp5_reference,
            "regulatory_clause_mapping": (
                self.regulatory_clause_mapping
            ),
            "recommendation": self.recommendation,
        }


@dataclass
class GapAnalysisReport:
    """
    Complete GAMP 5 gap analysis report for a vendor document.

    :requirement: URS-9.1 - System shall perform GAMP 5 gap analysis.
    """
    file_name: str
    title: str
    analyzed_at: str
    total_categories: int
    covered: int
    partial: int
    gaps: int
    findings: List[GapFinding] = field(default_factory=list)
    summary: str = ""
    limitations: List[str] = field(default_factory=list)
    requirement_mappings: List[Dict[str, Any]] = field(
        default_factory=list
    )

    def to_dict(self) -> Dict[str, Any]:
        """
        Convert report to dictionary.

        :return: Dictionary representation of the report.
        :requirement: URS-9.3 - System shall output gap analysis as JSON.
        """
        return {
            "file_name": self.file_name,
            "title": self.title,
            "analyzed_at": self.analyzed_at,
            "total_categories": self.total_categories,
            "covered": self.covered,
            "partial": self.partial,
            "gaps": self.gaps,
            "summary": self.summary,
            "limitations": self.limitations,
            "findings": [f.to_dict() for f in self.findings],
            "requirement_mappings": self.requirement_mappings,
        }

    def to_json(self) -> str:
        """
        Convert report to JSON string.

        :return: JSON string representation.
        :requirement: URS-9.3 - System shall output gap analysis as JSON.
        """
        return json.dumps(self.to_dict(), indent=2)

    def save(self, output_dir: str = OUTPUT_DIR) -> Path:
        """
        Save report to output/gap_analysis_report.json.

        :param output_dir: Directory to write the report file.
        :return: Path to the saved file.
        :requirement: URS-9.4 - System shall save gap analysis report.
        """
        out_path = Path(output_dir)
        out_path.mkdir(parents=True, exist_ok=True)
        file_path = out_path / "gap_analysis_report.json"
        file_path.write_text(self.to_json(), encoding="utf-8")
        return file_path


# ── IngestorAgent ─────────────────────────────────────────────────────
class IngestorAgent:
    """
    Reads vendor .docx and .pdf files and converts them to structured
    JSON for downstream consumption by the RequirementArchitect.

    :requirement: URS-8.1 - System shall ingest vendor documents.
    """

    # Coverage thresholds for semantic gap scoring
    COVERED_THRESHOLD = 0.50
    PARTIAL_THRESHOLD = 0.35

    def __init__(
        self,
        vendor_docs_dir: str = VENDOR_DOCS_DIR
    ):
        """
        Initialize the IngestorAgent.

        :param vendor_docs_dir: Path to the vendor documents directory.
        :requirement: URS-8.1 - System shall ingest vendor documents.
        """
        self._vendor_docs_dir = vendor_docs_dir
        self._validate_dependencies()

    # ── Dependency validation ─────────────────────────────────────────
    def _validate_dependencies(self) -> None:
        """
        Validate that at least one PDF parser and docx are available.

        :raises ImportError: If no PDF parser or python-docx is missing.
        :requirement: URS-8.5 - System shall validate dependencies.
        """
        if DocxDocument is None:
            raise ImportError(
                "python-docx is required. "
                "Install with: pip install python-docx"
            )
        if _pdfplumber is None and PdfReader is None:
            raise ImportError(
                "A PDF parser is required. "
                "Install with: pip install pdfplumber"
            )

    # ── Manifest (SHA-256 deduplication) ──────────────────────────────
    @staticmethod
    def _compute_document_hash(file_path: str) -> str:
        """
        Compute SHA-256 hash of a file for deduplication.

        :param file_path: Absolute path to the file.
        :return: Hex-encoded SHA-256 hash string.
        :requirement: URS-8.11 - System shall deduplicate documents.
        """
        h = hashlib.sha256()
        with open(file_path, "rb") as fh:
            for chunk in iter(lambda: fh.read(65536), b""):
                h.update(chunk)
        return h.hexdigest()

    @staticmethod
    def _load_manifest() -> Dict[str, Any]:
        """
        Load the ingest manifest from disk.

        :return: Manifest dict (empty if file absent).
        :requirement: URS-8.11 - System shall deduplicate documents.
        """
        if MANIFEST_PATH.exists():
            try:
                return json.loads(
                    MANIFEST_PATH.read_text(encoding="utf-8")
                )
            except Exception:
                pass
        return {}

    @staticmethod
    def _save_manifest_entry(
        doc_hash: str,
        name: str,
        sections: int,
        reqs: int,
        lims: int,
    ) -> None:
        """
        Save a manifest entry for a processed document.

        :param doc_hash: SHA-256 hash of the file.
        :param name: Original file name.
        :param sections: Number of sections extracted.
        :param reqs: Number of requirements extracted.
        :param lims: Number of limitations extracted.
        :requirement: URS-8.11 - System shall deduplicate documents.
        """
        MANIFEST_PATH.parent.mkdir(parents=True, exist_ok=True)
        manifest = IngestorAgent._load_manifest()
        manifest[doc_hash] = {
            "file_name": name,
            "first_ingested": datetime.now(
                timezone.utc
            ).isoformat(),
            "sections": sections,
            "requirements": reqs,
            "limitations": lims,
        }
        MANIFEST_PATH.write_text(
            json.dumps(manifest, indent=2),
            encoding="utf-8"
        )

    @staticmethod
    def _check_duplicate(
        doc_hash: str,
    ) -> Optional[Dict[str, Any]]:
        """
        Check if a document hash exists in the manifest.

        :param doc_hash: SHA-256 hash to look up.
        :return: Existing manifest entry dict or None.
        :requirement: URS-8.11 - System shall deduplicate documents.
        """
        manifest = IngestorAgent._load_manifest()
        return manifest.get(doc_hash)

    # ── Table helper ──────────────────────────────────────────────────
    @staticmethod
    def _table_to_markdown(
        rows: List[List[Optional[str]]]
    ) -> str:
        """
        Convert a 2-D list of cell strings to a Markdown table.

        :param rows: List of rows; each row is a list of cell strings
                     (None cells are treated as empty strings).
        :return: Markdown table string.
        :requirement: URS-8.12 - System shall extract table content.
        """
        if not rows:
            return ""

        # Normalise cells
        str_rows: List[List[str]] = [
            [
                (str(c) if c is not None else "").replace("|", "\\|")
                for c in row
            ]
            for row in rows
        ]

        # Pad columns to uniform width
        col_count = max(len(r) for r in str_rows)
        for row in str_rows:
            while len(row) < col_count:
                row.append("")

        col_widths = [
            max(len(str_rows[r][c]) for r in range(len(str_rows)))
            for c in range(col_count)
        ]

        def _fmt_row(cells: List[str]) -> str:
            parts = [
                cells[c].ljust(col_widths[c])
                for c in range(col_count)
            ]
            return "| " + " | ".join(parts) + " |"

        lines = [_fmt_row(str_rows[0])]
        sep = (
            "| "
            + " | ".join("-" * col_widths[c] for c in range(col_count))
            + " |"
        )
        lines.append(sep)
        for row in str_rows[1:]:
            lines.append(_fmt_row(row))
        return "\n".join(lines)

    # ── Heading heuristic ─────────────────────────────────────────────
    @staticmethod
    def _is_heading_line(
        text: str,
        font_size: float,
        median_size: float,
    ) -> bool:
        """
        Heuristic to decide whether a text line is a section heading.

        :param text: The line text.
        :param font_size: Font size of the line (0 if unknown).
        :param median_size: Median font size of the document.
        :return: True if the line looks like a heading.
        :requirement: URS-8.7 - System shall parse PDF documents.
        """
        t = text.strip()
        if not t:
            return False
        if font_size > 0 and median_size > 0:
            if font_size > median_size * 1.25:
                return True
        if len(t) <= 80 and t == t.upper() and len(t) > 3:
            return True
        if _HEADING_RE.match(t):
            return True
        return False

    # ── PDF parsing (pdfplumber primary) ──────────────────────────────
    def _parse_pdf_plumber(
        self, file_path: str
    ) -> IngestedDocument:
        """
        Parse a PDF using pdfplumber for structural/table awareness.

        :param file_path: Absolute path to the PDF file.
        :return: Structured IngestedDocument.
        :raises DocumentParseError: If the PDF cannot be read.
        :requirement: URS-8.7 - System shall parse PDF documents.
        """
        try:
            pdf = _pdfplumber.open(file_path)
        except Exception as exc:
            raise DocumentParseError(file_path, str(exc))

        sections: List[DocumentSection] = []
        section_idx = 0
        cur_heading = "Introduction"
        cur_lines: List[str] = []
        title: str = ""
        table_count = 0

        # Compute median font size from first 10 pages
        sample_sizes: List[float] = []
        for pg in pdf.pages[:10]:
            for ch in (pg.chars or []):
                sz = ch.get("size", 0)
                if sz:
                    sample_sizes.append(float(sz))
        median_size = (
            statistics.median(sample_sizes) if sample_sizes else 0.0
        )

        def _flush_body() -> None:
            nonlocal section_idx, title
            text = "\n".join(cur_lines).strip()
            if text:
                sections.append(DocumentSection(
                    heading=cur_heading,
                    content=text,
                    page_number=page_num,
                    section_index=section_idx,
                    section_type="body",
                ))
                if not title and cur_heading.lower() not in (
                    "introduction", ""
                ):
                    title = cur_heading
                section_idx += 1
            cur_lines.clear()

        for page_num_0, page in enumerate(pdf.pages):
            page_num = page_num_0 + 1

            # Find table bounding boxes to exclude from text chars
            tables = page.find_tables()
            table_bboxes = [
                tbl.bbox for tbl in tables
            ] if tables else []

            # Extract and emit tables as Markdown sections
            raw_tables = page.extract_tables() or []
            for tbl_rows in raw_tables:
                _flush_body()
                md = self._table_to_markdown(tbl_rows)
                if md:
                    sections.append(DocumentSection(
                        heading=f"Table (p.{page_num})",
                        content=md,
                        page_number=page_num,
                        section_index=section_idx,
                        section_type="table",
                    ))
                    section_idx += 1
                    table_count += 1

            # Build line map from page chars, skipping table regions
            def _in_table_bbox(top: float, bottom: float) -> bool:
                for (x0, y0, x1, y1) in table_bboxes:
                    if top >= y0 and bottom <= y1:
                        return True
                return False

            line_map: Dict[float, List[Dict[str, Any]]] = {}
            for ch in (page.chars or []):
                if _in_table_bbox(ch["top"], ch["bottom"]):
                    continue
                key = round(ch["top"], 1)
                line_map.setdefault(key, []).append(ch)

            for top_key in sorted(line_map.keys()):
                chars = sorted(line_map[top_key], key=lambda c: c["x0"])
                line_text = "".join(c["text"] for c in chars).strip()
                if not line_text:
                    continue
                font_size = float(
                    chars[0].get("size", 0) if chars else 0
                )

                if self._is_heading_line(
                    line_text, font_size, median_size
                ):
                    _flush_body()
                    cur_heading = line_text
                else:
                    cur_lines.append(line_text)

        # Flush last body section using last page_num
        if cur_lines:
            _flush_body()

        total_pages = len(pdf.pages)  # capture before close
        pdf.close()

        if not title:
            title = (
                Path(file_path).stem
                .replace("_", " ")
                .replace("-", " ")
            )

        full_text = "\n".join(s.content for s in sections)
        return IngestedDocument(
            file_name=Path(file_path).name,
            file_type="pdf",
            title=title,
            ingested_at=datetime.now(timezone.utc).isoformat(),
            total_pages=total_pages,
            sections=sections,
        )

    def _parse_pdf_fallback(
        self, file_path: str
    ) -> IngestedDocument:
        """
        Parse a PDF using PyPDF2 (fallback when pdfplumber unavailable).

        :param file_path: Absolute path to the PDF file.
        :return: Structured IngestedDocument.
        :raises DocumentParseError: If the PDF cannot be read.
        :requirement: URS-8.7 - System shall parse PDF documents.
        """
        try:
            reader = PdfReader(file_path)
        except Exception as exc:
            raise DocumentParseError(file_path, str(exc))

        sections: List[DocumentSection] = []
        section_idx = 0
        cur_heading = "Introduction"
        cur_lines: List[str] = []
        title = ""

        def _flush_body(page_num: int) -> None:
            nonlocal section_idx, title
            text = "\n".join(cur_lines).strip()
            if text:
                sections.append(DocumentSection(
                    heading=cur_heading,
                    content=text,
                    page_number=page_num,
                    section_index=section_idx,
                    section_type="body",
                ))
                if not title and cur_heading.lower() != "introduction":
                    title = cur_heading
                section_idx += 1
            cur_lines.clear()

        for page_num_0, page in enumerate(reader.pages):
            page_num = page_num_0 + 1
            page_text = page.extract_text() or ""
            for line in page_text.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                if self._is_heading_line(stripped, 12.0, 12.0):
                    _flush_body(page_num)
                    cur_heading = stripped
                else:
                    cur_lines.append(stripped)

            _flush_body(page_num)

        if not title:
            title = (
                Path(file_path).stem
                .replace("_", " ")
                .replace("-", " ")
            )

        return IngestedDocument(
            file_name=Path(file_path).name,
            file_type="pdf",
            title=title,
            ingested_at=datetime.now(timezone.utc).isoformat(),
            total_pages=len(reader.pages),
            sections=sections,
        )

    # ── DOCX parsing (enhanced with table extraction) ─────────────────
    def _parse_docx(
        self, file_path: str
    ) -> IngestedDocument:
        """
        Parse a .docx file, extracting paragraphs and tables.

        :param file_path: Absolute path to the .docx file.
        :return: Structured IngestedDocument.
        :raises DocumentParseError: If the file cannot be read.
        :requirement: URS-8.8 - System shall parse DOCX documents.
        """
        try:
            doc = DocxDocument(file_path)
        except Exception as exc:
            raise DocumentParseError(file_path, str(exc))

        sections: List[DocumentSection] = []
        section_idx = 0
        cur_heading = "Introduction"
        cur_lines: List[str] = []
        page_num = 1

        def _flush_body() -> None:
            nonlocal section_idx
            text = "\n".join(cur_lines).strip()
            if text:
                sections.append(DocumentSection(
                    heading=cur_heading,
                    content=text,
                    page_number=page_num,
                    section_index=section_idx,
                    section_type="body",
                ))
                section_idx += 1
            cur_lines.clear()

        if _DOCX_HAS_XML and _docx_qn is not None:
            # Enhanced path: iterate XML body children
            for child in doc.element.body:
                tag = child.tag

                if tag == _docx_qn("w:p"):
                    # Paragraph
                    from docx.text.paragraph import Paragraph
                    para = Paragraph(child, doc)
                    text = para.text.strip()
                    if not text:
                        continue
                    style_name = (
                        para.style.name or ""
                    ).lower()
                    if "heading" in style_name:
                        _flush_body()
                        cur_heading = text
                    else:
                        cur_lines.append(text)

                elif tag == _docx_qn("w:tbl"):
                    # Table
                    _flush_body()
                    from docx.table import Table
                    tbl = Table(child, doc)
                    rows = [
                        [
                            cell.text.strip()
                            for cell in row.cells
                        ]
                        for row in tbl.rows
                    ]
                    md = self._table_to_markdown(rows)
                    if md:
                        sections.append(DocumentSection(
                            heading=f"Table (p.{page_num})",
                            content=md,
                            page_number=page_num,
                            section_index=section_idx,
                            section_type="table",
                        ))
                        section_idx += 1
        else:
            # Fallback: paragraph-only loop
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = (para.style.name or "").lower()
                if "heading" in style_name:
                    _flush_body()
                    cur_heading = text
                else:
                    cur_lines.append(text)

        _flush_body()

        # Title from core properties
        title = (
            Path(file_path).stem
            .replace("_", " ")
            .replace("-", " ")
        )
        try:
            if doc.core_properties.title:
                title = doc.core_properties.title
        except Exception:
            pass

        return IngestedDocument(
            file_name=Path(file_path).name,
            file_type="docx",
            title=title,
            ingested_at=datetime.now(timezone.utc).isoformat(),
            total_pages=len(sections),
            sections=sections,
        )

    # ── Requirement/limitation extraction ─────────────────────────────
    def _extract_requirements_from_text(
        self, text: str
    ) -> List[str]:
        """
        Extract requirement-like statements via regex.

        :param text: Raw text to scan.
        :return: List of extracted requirement strings.
        :requirement: URS-8.6 - System shall extract requirements.
        """
        requirements: List[str] = []
        seen: set = set()

        shall_pattern = re.compile(
            r'[^.]*\b(?:shall|must|is required to|'
            r'should|needs? to)\b[^.]*\.',
            re.IGNORECASE
        )
        for match in shall_pattern.finditer(text):
            req = match.group(0).strip()
            req_norm = req.lower()
            if req_norm not in seen and len(req) > 20:
                seen.add(req_norm)
                requirements.append(req)

        numbered_pattern = re.compile(
            r'(?:^|\n)\s*(?:\d+[\.\)]\s*\d*[\.\)]?\s*|'
            r'[a-z][\.\)]\s*)(.{20,}?)(?:\n|$)',
            re.IGNORECASE
        )
        for match in numbered_pattern.finditer(text):
            req = match.group(1).strip().rstrip('.')
            req_norm = req.lower()
            if req_norm not in seen and len(req) > 20:
                seen.add(req_norm)
                requirements.append(req)

        return requirements

    def _extract_limitations_regex(
        self, text: str
    ) -> List[str]:
        """
        Extract limitation statements via regex patterns.

        :param text: Raw text to scan.
        :return: List of limitation strings.
        :requirement: URS-8.13 - System shall extract limitations.
        """
        limitations: List[str] = []
        seen: set = set()
        for match in _LIMITATION_PATTERNS.finditer(text):
            lim = match.group(0).strip()
            lim_norm = lim.lower()
            if lim_norm not in seen and len(lim) > 15:
                seen.add(lim_norm)
                limitations.append(lim)
        return limitations

    def _extract_requirements_llm(
        self, text: str
    ) -> Tuple[List[str], List[str]]:
        """
        Use gpt-4o-mini to extract requirements and limitations.

        Processes text in 4000-char chunks. Falls back to ([], [])
        on any error (network, API key missing, etc.).

        :param text: Full document text.
        :return: Tuple of (requirements, limitations) lists.
        :requirement: URS-8.6 - System shall extract requirements.
        :requirement: URS-8.13 - System shall extract limitations.
        """
        if _OpenAI is None:
            return [], []

        openai_api_key = os.getenv("OPENAI_API_KEY")
        if not openai_api_key:
            return [], []

        client = _OpenAI(api_key=openai_api_key)
        all_reqs: List[str] = []
        all_lims: List[str] = []
        seen_reqs: set = set()
        seen_lims: set = set()

        chunk_size = 4000
        for i in range(0, len(text), chunk_size):
            chunk = text[i:i + chunk_size]
            prompt = (
                "Extract explicit AND implied requirements and "
                "limitations from the following text. "
                "Return ONLY valid JSON with this exact structure:\n"
                '{"requirements": ["..."], "limitations": ["..."]}\n'
                "Requirements are statements describing what the "
                "system shall/must/should do. "
                "Limitations are prohibited actions, out-of-scope "
                "items, or constraints (shall not, cannot, etc.).\n\n"
                f"TEXT:\n{chunk}"
            )
            try:
                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    temperature=0,
                    max_tokens=1200,
                    messages=[
                        {
                            "role": "system",
                            "content": (
                                "You are a regulatory requirements "
                                "analyst. Output only valid JSON."
                            ),
                        },
                        {"role": "user", "content": prompt},
                    ],
                )
                raw = resp.choices[0].message.content or ""
                # Strip markdown fences
                raw = re.sub(
                    r'^```(?:json)?\s*|\s*```$', '', raw.strip()
                )
                parsed = json.loads(raw)
                for r in parsed.get("requirements", []):
                    r_norm = r.strip().lower()
                    if r_norm not in seen_reqs and len(r.strip()) > 10:
                        seen_reqs.add(r_norm)
                        all_reqs.append(r.strip())
                for lim in parsed.get("limitations", []):
                    lim_norm = lim.strip().lower()
                    if (
                        lim_norm not in seen_lims
                        and len(lim.strip()) > 10
                    ):
                        seen_lims.add(lim_norm)
                        all_lims.append(lim.strip())
            except Exception:
                pass  # Non-critical; regex fallback still runs

        return all_reqs, all_lims

    # ── Audit helper ──────────────────────────────────────────────────
    def _log_audit_event(
        self, action: str, details: Dict[str, Any]
    ) -> None:
        """
        Log an audit event for document ingestion.

        :param action: The audit action name.
        :param details: Additional details for the audit record.
        :requirement: URS-2.1 - System shall maintain audit trail.
        """
        audit_logger.info(
            action,
            extra={
                "user_id": "SYSTEM",
                "timestamp": datetime.now(timezone.utc).isoformat(),
                "action": action,
                "details": details,
            }
        )

    # ── Public API ────────────────────────────────────────────────────
    def ingest_file(
        self, file_path: str
    ) -> IngestedDocument:
        """
        Ingest a single vendor document and return structured JSON.

        Order of operations:
        1. Resolve path and validate extension
        2. Compute SHA-256 → check manifest (log if duplicate)
        3. Parse PDF (pdfplumber → fallback) or DOCX
        4. Assign document_hash to result
        5. Build full_text from sections
        6. LLM extraction → (llm_reqs, llm_lims)
        7. Regex extraction → (regex_reqs, regex_lims)
        8. Merge + deduplicate both lists
        9. Assign result.requirements and result.limitations
        10. Save manifest entry
        11. Log audit event

        :param file_path: Path to the .docx or .pdf file.
        :return: IngestedDocument with sections, requirements,
                 limitations, and document_hash.
        :raises UnsupportedFileTypeError: If the file type is invalid.
        :raises DocumentParseError: If the file cannot be parsed.
        :requirement: URS-8.1 - System shall ingest vendor documents.
        """
        # Step 1: Resolve path
        path = Path(file_path)
        if not path.is_absolute():
            path = Path(self._vendor_docs_dir) / path
        file_path = str(path)

        if not path.exists():
            raise DocumentParseError(file_path, "File does not exist")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(file_path)

        self._log_audit_event(
            "DOCUMENT_INGESTION_STARTED",
            {"file": file_path, "type": ext}
        )

        # Step 2: SHA-256 deduplication
        doc_hash = self._compute_document_hash(file_path)
        existing = self._check_duplicate(doc_hash)
        if existing:
            _log_integrity_event(
                agent_name="IngestorAgent",
                action="DOCUMENT_DUPLICATE_DETECTED",
                decision_logic=(
                    f"Duplicate detected: {path.name} "
                    f"(hash={doc_hash[:12]}…); "
                    f"previously ingested as "
                    f"{existing.get('file_name', 'unknown')}"
                ),
            )

        # Step 3: Parse
        if ext == ".pdf":
            if _pdfplumber is not None:
                result = self._parse_pdf_plumber(file_path)
            else:
                result = self._parse_pdf_fallback(file_path)
        else:
            result = self._parse_docx(file_path)

        # Step 4: Assign hash
        result.document_hash = doc_hash

        # Step 5: Build full text
        full_text = "\n".join(s.content for s in result.sections)

        # Step 6: LLM extraction
        llm_reqs, llm_lims = self._extract_requirements_llm(full_text)
        llm_used = bool(llm_reqs or llm_lims)

        # Step 7: Regex extraction
        regex_reqs = self._extract_requirements_from_text(full_text)
        regex_lims = self._extract_limitations_regex(full_text)

        # Step 8: Merge + deduplicate (LLM first)
        seen_reqs: set = set()
        merged_reqs: List[str] = []
        for r in llm_reqs + regex_reqs:
            r_norm = r.strip().lower()
            if r_norm not in seen_reqs:
                seen_reqs.add(r_norm)
                merged_reqs.append(r.strip())

        seen_lims: set = set()
        merged_lims: List[str] = []
        for lim in llm_lims + regex_lims:
            lim_norm = lim.strip().lower()
            if lim_norm not in seen_lims:
                seen_lims.add(lim_norm)
                merged_lims.append(lim.strip())

        # Step 9: Assign
        result.requirements = merged_reqs
        result.limitations = merged_lims

        # Step 10: Save manifest
        table_count = sum(
            1 for s in result.sections if s.section_type == "table"
        )
        self._save_manifest_entry(
            doc_hash=doc_hash,
            name=path.name,
            sections=len(result.sections),
            reqs=len(merged_reqs),
            lims=len(merged_lims),
        )

        # Step 11: Audit
        self._log_audit_event(
            "DOCUMENT_INGESTION_COMPLETED",
            {
                "file": file_path,
                "sections": len(result.sections),
                "tables": table_count,
                "requirements_found": len(merged_reqs),
                "limitations_found": len(merged_lims),
                "llm_extraction": llm_used,
                "document_hash": doc_hash[:16] + "…",
            }
        )

        decision_logic = (
            f"Parsed {result.file_name} "
            f"({ext.lstrip('.').upper()}, "
            f"{result.total_pages} pages); "
            f"extracted {len(result.sections)} sections "
            f"({table_count} tables), "
            f"{len(merged_reqs)} requirements, "
            f"{len(merged_lims)} limitations; "
            f"LLM extraction: {'yes' if llm_used else 'no'}"
        )
        _log_integrity_event(
            agent_name="IngestorAgent",
            action="DOCUMENT_INGESTED",
            decision_logic=decision_logic,
        )

        return result

    def ingest_all(self) -> List[IngestedDocument]:
        """
        Ingest all supported documents from the vendor_docs directory.

        :return: List of IngestedDocument objects.
        :raises FileNotFoundError: If the vendor_docs directory is absent.
        :requirement: URS-8.9 - System shall batch-ingest documents.
        """
        docs_dir = Path(self._vendor_docs_dir)
        if not docs_dir.exists():
            raise FileNotFoundError(
                f"Vendor docs directory not found: "
                f"{self._vendor_docs_dir}"
            )

        results: List[IngestedDocument] = []
        all_files = sorted(
            p for p in docs_dir.iterdir()
            if p.suffix.lower() in SUPPORTED_EXTENSIONS
        )

        for file_path in all_files:
            try:
                doc = self.ingest_file(str(file_path))
                results.append(doc)
            except IngestorError as exc:
                audit_logger.warning(
                    "DOCUMENT_INGESTION_FAILED",
                    extra={
                        "user_id": "SYSTEM",
                        "timestamp": datetime.now(
                            timezone.utc
                        ).isoformat(),
                        "action": "DOCUMENT_INGESTION_FAILED",
                        "details": {
                            "file": str(file_path),
                            "error": str(exc),
                        },
                    }
                )
                _log_integrity_event(
                    agent_name="IngestorAgent",
                    action="DOCUMENT_INGESTION_FAILED",
                    decision_logic=(
                        f"Failed to ingest {file_path.name}: {exc}"
                    ),
                )

        succeeded = len(results)
        failed = len(all_files) - succeeded
        _log_integrity_event(
            agent_name="IngestorAgent",
            action="BATCH_INGESTION_COMPLETED",
            decision_logic=(
                f"Batch-ingested documents from "
                f"{self._vendor_docs_dir}; "
                f"{succeeded} succeeded, {failed} failed"
            ),
        )
        return results

    def get_requirements_for_architect(
        self, file_path: str
    ) -> List[str]:
        """
        Ingest a document and return its requirements list.

        :param file_path: Path to the vendor document.
        :return: List of requirement strings.
        :requirement: URS-8.10 - System shall feed RequirementArchitect.
        """
        doc = self.ingest_file(file_path)
        return doc.requirements

    # ── Keyword evidence (kept for legacy callers) ────────────────────
    def _find_keyword_evidence(
        self,
        text: str,
        keywords: List[str],
    ) -> str:
        """
        Search text for keywords and return an excerpt as evidence.

        :param text: Full document text to search.
        :param keywords: List of indicator keywords for a category.
        :return: ~200-char excerpt around the first match, or empty.
        :requirement: URS-9.5 - System shall extract vendor evidence.
        """
        text_lower = text.lower()
        for keyword in keywords:
            pos = text_lower.find(keyword)
            if pos != -1:
                start = max(0, pos - 40)
                end = min(
                    len(text),
                    pos + len(keyword) + EVIDENCE_EXCERPT_LENGTH
                )
                excerpt = text[start:end].strip()
                if start > 0:
                    excerpt = "..." + excerpt
                if end < len(text):
                    excerpt = excerpt + "..."
                return excerpt
        return ""

    # ── Gap analysis (semantic) ───────────────────────────────────────
    def analyze_gaps(
        self,
        file_path: str,
        pinecone_api_key: Optional[str] = None,
        openai_api_key: Optional[str] = None,
    ) -> GapAnalysisReport:
        """
        Perform semantic GAMP 5 gap analysis on a vendor document.

        Uses RequirementArchitect.search() similarity scores instead of
        keyword matching. Coverage thresholds:
          - score >= 0.50 → "covered"
          - score >= 0.35 → "partial"
          - below         → "gap"

        :param file_path: Path to the vendor .docx or .pdf file.
        :param pinecone_api_key: Pinecone API key (defaults to env).
        :param openai_api_key: OpenAI API key (defaults to env).
        :return: GapAnalysisReport with semantic findings and summary.
        :raises IngestorError: If the document cannot be ingested.
        :requirement: URS-9.1 - System shall perform GAMP 5 gap analysis.
        """
        from Agents.requirement_architect import RequirementArchitect

        self._log_audit_event(
            "GAP_ANALYSIS_STARTED", {"file": file_path}
        )

        # Step 1: Ingest
        doc = self.ingest_file(file_path)
        full_text = "\n".join(
            section.content for section in doc.sections
        )

        # Step 2: Architect for Pinecone queries
        architect = RequirementArchitect(
            pinecone_api_key=pinecone_api_key,
            openai_api_key=openai_api_key,
        )

        # Step 3: Semantic category analysis
        findings: List[GapFinding] = []
        covered_names: List[str] = []
        partial_names: List[str] = []
        gap_names: List[str] = []

        for cat in GAMP5_CATEGORIES:
            category_name: str = cat["category"]
            keywords: List[str] = cat["keywords"]

            # Build query from vendor excerpt or fallback to category query
            vendor_excerpt = self._find_keyword_evidence(
                full_text, keywords
            )
            query = (
                vendor_excerpt[:500]
                if vendor_excerpt
                else cat["query"]
            )

            # Query Pinecone for top-3 semantic matches
            best_score = 0.0
            gamp5_ref = "GAMP 5 reference unavailable"
            clause_mapping: List[Dict[str, Any]] = []

            try:
                search_resp = architect.search(
                    query=query,
                    top_k=3,
                    min_score=0.0,
                )
                results = search_resp.results if search_resp else []
                if results:
                    best_score = float(
                        results[0].similarity_score or 0.0
                    )
                    # Format primary reference
                    r0 = results[0]
                    src = r0.source_document or "GAMP 5"
                    pg = r0.page_number or 0
                    ver = r0.reg_version or ""
                    txt = (r0.text or "")[:200]
                    if ver:
                        gamp5_ref = (
                            f"Per {src} [{ver}] (p.{pg}): {txt}..."
                        )
                    else:
                        gamp5_ref = (
                            f"Per {src} (p.{pg}): {txt}..."
                        )
                    # Build regulatory clause mapping (top-3)
                    for rank_0, r in enumerate(results, start=1):
                        clause_mapping.append({
                            "rank": rank_0,
                            "source": r.source_document or "",
                            "page": r.page_number or 0,
                            "reg_version": r.reg_version or "",
                            "text_excerpt": (
                                (r.text or "")[:300]
                            ),
                            "similarity_score": round(
                                float(r.similarity_score or 0.0), 4
                            ),
                        })
            except Exception:
                gamp5_ref = "GAMP 5 reference unavailable"

            # Determine status from similarity score
            if best_score >= self.COVERED_THRESHOLD:
                status = "covered"
                covered_names.append(category_name)
                recommendation = ""
            elif best_score >= self.PARTIAL_THRESHOLD:
                status = "partial"
                partial_names.append(category_name)
                recommendation = (
                    f"Vendor document partially addresses "
                    f"{category_name.lower()}. "
                    f"Consider expanding coverage."
                )
            else:
                status = "gap"
                gap_names.append(category_name)
                recommendation = (
                    f"Vendor should provide documentation "
                    f"addressing {category_name.lower()}."
                )

            findings.append(GapFinding(
                category=category_name,
                status=status,
                similarity_score=best_score,
                vendor_evidence=vendor_excerpt,
                gamp5_reference=gamp5_ref,
                regulatory_clause_mapping=clause_mapping,
                recommendation=recommendation,
            ))

        # Step 4: Per-requirement regulatory mappings (first 20)
        requirement_mappings: List[Dict[str, Any]] = []
        for req in doc.requirements[:20]:
            req_clauses: List[Dict[str, Any]] = []
            try:
                resp = architect.search(
                    query=req[:500], top_k=3, min_score=0.0
                )
                for rank_0, r in enumerate(
                    resp.results if resp else [], start=1
                ):
                    req_clauses.append({
                        "rank": rank_0,
                        "source": r.source_document or "",
                        "page": r.page_number or 0,
                        "reg_version": r.reg_version or "",
                        "text_excerpt": (r.text or "")[:300],
                        "similarity_score": round(
                            float(r.similarity_score or 0.0), 4
                        ),
                    })
            except Exception:
                pass
            requirement_mappings.append({
                "requirement": req,
                "regulatory_clauses": req_clauses,
            })

        # Step 5: Summary
        total = len(GAMP5_CATEGORIES)
        covered_count = len(covered_names)
        partial_count = len(partial_names)
        gap_count = len(gap_names)

        parts: List[str] = []
        if covered_names:
            parts.append(f"covers {', '.join(covered_names)}")
        if partial_names:
            parts.append(
                f"partially covers {', '.join(partial_names)}"
            )
        if gap_names:
            parts.append(f"is missing {', '.join(gap_names)}")

        if parts:
            summary = f"Vendor document {'; '.join(parts)}."
        else:
            summary = (
                "Vendor document covers all assessed "
                "GAMP 5 lifecycle categories."
            )

        report = GapAnalysisReport(
            file_name=doc.file_name,
            title=doc.title,
            analyzed_at=datetime.now(timezone.utc).isoformat(),
            total_categories=total,
            covered=covered_count,
            partial=partial_count,
            gaps=gap_count,
            findings=findings,
            summary=summary,
            limitations=doc.limitations,
            requirement_mappings=requirement_mappings,
        )

        # Step 6: Save
        saved_path = report.save()

        self._log_audit_event(
            "GAP_ANALYSIS_COMPLETED",
            {
                "file": file_path,
                "total_categories": total,
                "covered": covered_count,
                "partial": partial_count,
                "gaps": gap_count,
                "report_path": str(saved_path),
            }
        )

        if gap_names:
            gap_detail = f"Gaps: {', '.join(gap_names)}"
        else:
            gap_detail = "No gaps identified"

        _log_integrity_event(
            agent_name="IngestorAgent",
            action="GAP_ANALYSIS_COMPLETED",
            decision_logic=(
                f"Analyzed {doc.file_name} against "
                f"{total} GAMP 5 categories using semantic "
                f"similarity; {covered_count} covered, "
                f"{partial_count} partial, "
                f"{gap_count} gaps. {gap_detail}"
            ),
        )

        return report
